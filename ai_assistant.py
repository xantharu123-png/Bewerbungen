import anthropic
import PyPDF2
import io
import os
from pathlib import Path
from datetime import datetime

# API Key aus Environment Variable oder Streamlit Secrets
def _load_api_key() -> str:
    """Load API key from environment or Streamlit secrets."""
    # 1. Environment variable (lokal oder Streamlit Cloud)
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if key:
        return key
    # 2. Streamlit secrets (secrets.toml oder Cloud Secrets)
    try:
        import streamlit as st
        key = st.secrets.get("ANTHROPIC_API_KEY", "")
        if key:
            return key
    except Exception:
        pass
    return ""

ANTHROPIC_API_KEY = _load_api_key()

def _get_client() -> anthropic.Anthropic:
    """Create Anthropic client with the configured API key."""
    return anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"[Fehler beim PDF-Lesen: {e}]"

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes, preserving bullet points and structure."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(docx_bytes))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            # Detect bullet/list paragraphs by style or numbering
            is_bullet = False
            style_name = (para.style.name or "").lower()
            if "list" in style_name or "bullet" in style_name or "aufz" in style_name:
                is_bullet = True
            # Also check if paragraph has numbering format (Word list items)
            if hasattr(para, '_element'):
                numPr = para._element.find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr'
                )
                if numPr is not None:
                    is_bullet = True
            # Check for manual bullets
            if text.startswith(('- ', '• ', '– ', '* ')):
                is_bullet = True

            if is_bullet and not text.startswith('- '):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        return "\n".join(lines)
    except Exception as e:
        return f"[Fehler beim DOCX-Lesen: {e}]"

def extract_company_details(company: str, job_description: str) -> dict:
    """Extract company address and contact person from job description using Claude.

    Returns dict with keys: contact_person, company_address
    """
    clean_company = _sanitize_company(company)
    if not clean_company:
        return {"contact_person": "", "company_address": ""}

    try:
        client = _get_client()
        prompt = f"""Aus der folgenden Stellenanzeige für "{clean_company}": extrahiere bitte:

1. Die **Kontaktperson** (Name der zuständigen Person für Bewerbungen, z.B. "Frau Anna Müller").
   - Suche nach Formulierungen wie "Kontakt:", "Ansprechperson:", "Fragen beantwortet", "bewirb dich bei" etc.
   - Wenn KEINE Kontaktperson im Inserat steht, gib einen leeren String zurück.

2. Die **Firmenadresse** des Unternehmens "{clean_company}" in der Schweiz.
   - Suche zuerst im Inserat nach einer Adresse (Strasse + PLZ + Ort).
   - Wenn KEINE Adresse im Inserat steht, verwende dein Wissen über den Schweizer Hauptsitz/Standort von "{clean_company}". Die Firmenadressen grosser Schweizer Unternehmen sind öffentlich bekannt.
   - Format IMMER: "Strasse Nr\\nPLZ Ort" (z.B. "Schiffbaustrasse 2\\n8005 Zürich")
   - Wenn du die Adresse weder im Inserat findest noch sicher kennst, gib einen leeren String zurück.

Stellenanzeige:
{job_description[:2500]}

Antworte NUR mit einem JSON-Objekt:
{{"contact_person": "Frau/Herr Vorname Nachname oder leer", "company_address": "Strasse Nr\\nPLZ Ort oder leer"}}"""

        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=200,
            messages=[{"role": "user", "content": prompt}]
        )

        import json as _json
        text = message.content[0].text.strip()
        json_match = __import__('re').search(r'\{.*\}', text, __import__('re').DOTALL)
        if json_match:
            result = _json.loads(json_match.group(0))
            return {
                "contact_person": result.get("contact_person", ""),
                "company_address": result.get("company_address", ""),
            }
    except Exception as e:
        print(f"[AI] extract_company_details error: {e}")

    return {"contact_person": "", "company_address": ""}


def generate_cover_letter(
    job_title: str,
    company: str,
    job_description: str,
    cv_text: str,
    existing_letter: str = "",
    language: str = "de",
    api_key: str = "",
    contact_person: str = "",
    company_address: str = "",
) -> str:
    """Generate a tailored cover letter using Claude API.

    Returns ONLY the body text: greeting → paragraphs with bullet list → closing → signature.
    """

    client = _get_client()

    lang_instruction = "auf Deutsch" if language == "de" else "in English"

    # Determine greeting
    if contact_person:
        name_parts = contact_person.strip().split()
        first_name = name_parts[0] if name_parts else ""
        if first_name.lower() in ["frau", "fr."]:
            greeting = f"Sehr geehrte Frau {' '.join(name_parts[1:])}"
        elif first_name.lower() in ["herr", "hr."]:
            greeting = f"Sehr geehrter Herr {' '.join(name_parts[1:])}"
        else:
            greeting = f"Sehr geehrte/r {contact_person}"
    else:
        greeting = "Sehr geehrte Damen und Herren"

    # ── Two modes: TEMPLATE mode (if existing letter) or FREE mode ──
    if existing_letter and len(existing_letter.strip()) > 100:
        # TEMPLATE MODE: Use existing letter as exact template
        prompt = f"""Du erhältst eine VORLAGE eines Bewerbungsschreibens und eine neue Stellenanzeige.

**DEINE AUFGABE:** Passe die Vorlage auf die neue Stelle an. Behalte dabei:
- EXAKT die gleiche Struktur und Anzahl Absätze
- EXAKT den gleichen Schreibstil und Tonfall
- EXAKT die gleiche Länge (ungefähr gleich viele Sätze pro Absatz)
- Die gleiche Art, wie der Bewerber sich vorstellt und argumentiert
- Falls die Vorlage eine Aufzählung mit "- " hat, behalte diese Aufzählung bei

Ändere NUR die stellenspezifischen Inhalte:
- Bezug zur neuen Stelle statt zur alten
- Relevante Erfahrungen/Kompetenzen für DIESE Stelle hervorheben
- Firmenname und Stellentitel anpassen
- Anrede anpassen
- **WICHTIG:** Lies die Stellenanzeige genau und identifiziere 1-2 spezifische Anforderungen oder Aufgaben, die der Bewerber laut CV erfüllt. Baue diese als zusätzliche Aufzählungspunkte ein (z.B. "Erfahrung in Pre-Sales Aktivitäten und Kunden-Workshops" oder "Kenntnisse in Schnittstellen-Entwicklung und Datenbanktechnologien"). So erkennt sich das Unternehmen im Anschreiben wieder.

**VORLAGE (diese Struktur und Stil EXAKT beibehalten):**
---
{existing_letter[:3000]}
---

**NEUE STELLE:**
Stelle: {job_title}
Unternehmen: {company}
Beschreibung: {job_description[:2000]}

**Lebenslauf (für relevante Details):**
{cv_text[:1500]}

**AUSGABE-REGELN:**
- Beginne mit: "{greeting}"
- Ende mit: "Freundliche Grüsse" und dann "Miroslav Mikulic" auf einer neuen Zeile
- KEIN Briefkopf, KEINE Adresse, KEIN Datum, KEIN Betreff — nur der Brieftext
- Aufzählungspunkte mit "- " am Zeilenanfang kennzeichnen
- Gib NUR den fertigen Brieftext zurück, keine Kommentare"""

    else:
        # FREE MODE: Generate from scratch — Swiss style with bullet list
        prompt = f"""Schreibe ein professionelles Schweizer Bewerbungsschreiben {lang_instruction}.

**Stellenanzeige:**
Stelle: {job_title}
Unternehmen: {company}
Beschreibung: {job_description[:2000]}

**Lebenslauf des Bewerbers:**
{cv_text[:2000]}

**WICHTIG — Gib NUR den Brieftext zurück, EXAKT in diesem Format:**

{greeting}

[Einleitung: 2-3 Sätze, warum diese Stelle mich anspricht und wie ich darauf aufmerksam wurde]

[Überleitung: 1-2 Sätze, dann "Was kann ich Ihrem Unternehmen bieten:" gefolgt von einer Aufzählung]

- [Stärke 1 mit konkretem Bezug zur Stelle]
- [Stärke 2 mit konkretem Bezug zum Lebenslauf]
- [Stärke 3 mit relevanter Erfahrung]
- [Stärke 4 falls passend]
- [INSERATS-MATCH: 1-2 spezifische Anforderungen aus der Stellenanzeige, die der Bewerber laut CV erfüllt — z.B. "Erfahrung in Pre-Sales und Kunden-Workshops durch langjährige Tätigkeit als Senior Consultant"]

[Schluss: 2-3 Sätze mit Motivation und Gesprächswunsch]

Freundliche Grüsse

Miroslav Mikulic

**REGELN:**
- Schreibe NUR die Absätze zwischen Anrede und Gruss — KEIN Briefkopf, KEINE Adresse, KEIN Datum, KEIN Betreff
- Beginne direkt mit "{greeting}"
- Ende mit "Freundliche Grüsse" und dann "Miroslav Mikulic"
- Die Aufzählung mit "- " am Zeilenanfang, 4-6 Punkte
- Professioneller, selbstbewusster Ton — schweizerisch
- Keine Floskeln, keine Superlative
- Aufzählungspunkte sollen konkret und auf die Stelle bezogen sein
- **WICHTIG:** Lies die Stellenanzeige genau und identifiziere 1-2 spezifische Anforderungen oder Aufgaben, die der Bewerber laut CV erfüllt. Baue diese als eigene Aufzählungspunkte ein, damit sich das Unternehmen im Anschreiben wiedererkennt (z.B. "Erfahrung in Pre-Sales Aktivitäten und Kunden-Workshops" oder "Hands-on-Mentalität durch eigenständige Projektführung bei BaseNet und Eurofins")"""

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    return message.content[0].text


def _sanitize_company(company: str) -> str:
    """Clean company name: remove scraped garbage, keep only the actual company name."""
    if not company:
        return ""
    # If company contains line breaks, take only the first line
    company = company.split('\n')[0].strip()

    # Always try to split on common scraper artifacts
    scraper_artifacts = ['Speichern', 'Bewerben', 'Gestern', 'Heute', 'Einfach bewerben']
    for artifact in scraper_artifacts:
        if artifact in company:
            company = company.split(artifact)[0].strip()

    # Split on ® — company name is before it, rest is usually location/garbage
    if '®' in company:
        parts = company.split('®')
        if len(parts[0]) > 2:
            after = '®'.join(parts[1:]).strip()
            # If after ® there's short text like " AG" keep it, otherwise drop
            if after and (after.startswith(' AG') or after.startswith(' GmbH') or after.startswith(' SA') or len(after) < 5):
                company = parts[0].strip() + '®' + after
            else:
                company = parts[0].strip()

    # If still too long, it's scraped garbage
    if len(company) > 60:
        # Try common separators
        for sep in [' - ', '|', '·', '  ']:
            if sep in company:
                company = company.split(sep)[0].strip()
                break
        if len(company) > 60:
            company = company[:50].strip()

    # Remove trailing punctuation artifacts
    company = company.rstrip('|·-– ')
    return company


def generate_cover_letter_pdf(
    letter_text: str,
    job_title: str = "",
    company: str = "",
    contact_person: str = "",
    company_address: str = "",
) -> bytes:
    """Generate PDF by filling the user's DOCX template and converting via LibreOffice.

    This guarantees pixel-perfect output matching Bewerbungsschreiben_Mikulic.docx
    because we modify the actual template (fonts, tabs, spacing all preserved)
    and let LibreOffice handle the PDF rendering.

    DOCX structure (P0-P31):
      P0:  Miroslav Mikulic          P1: Street  P2: City
      P3:  Phone [TAB] Company       P4: [TAB] Addr1  P5: [TAB] Addr2
      P6:  (empty)                   P7: [TAB] Date
      P8-9: (empty)                  P10: Betreff (bold)
      P11: (empty)                   P12: Sehr geehrte...
      P13: (empty)                   P14: Intro text
      P15: (empty)                   P16: Was kann ich...
      P17: (empty)                   P18-23: Bullets
      P24-25: (empty)               P26: Closing text
      P27: (empty)                   P28: Ich freue mich...
      P29: (empty)                   P30: Freundliche Grüsse
      P31: Miroslav Mikulic
    """
    import subprocess
    import tempfile
    from pathlib import Path as _Path

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        # Fallback: if python-docx not available, use basic reportlab
        return _generate_cover_letter_pdf_fallback(
            letter_text, job_title, company, contact_person, company_address
        )

    # ── Locate the DOCX template ──
    # First: try to restore from database (base64-embedded)
    try:
        from database import get_document, DOCS_PATH
        doc_info = get_document("cover_letter")
        if doc_info:
            db_path = _Path(doc_info.get("path", ""))
            if db_path.exists() and db_path.stat().st_size > 0:
                print(f"[PDF] Template from DB: {db_path}")
    except Exception as e:
        print(f"[PDF] DB template lookup failed: {e}")

    template_candidates = [
        _Path("data/documents/Bewerbungsschreiben_Mikulic.docx"),
        *list(_Path("data/documents").glob("Bewerbungsschreiben*.docx")),
        *list(_Path("data/documents").glob("*.docx")),
        _Path("Bewerbungsschreiben_Mikulic.docx"),
    ]

    template_path = None
    for p in template_candidates:
        if p.exists() and p.stat().st_size > 0:
            template_path = p
            break

    if not template_path:
        print("[PDF] No DOCX template found — using reportlab fallback")
        return _generate_cover_letter_pdf_fallback(
            letter_text, job_title, company, contact_person, company_address
        )

    # ── Parse the AI letter text into parts ──
    clean_company = _sanitize_company(company)
    body_text = letter_text.strip()

    # Strip header/address lines the AI may have prepended
    lines = body_text.split("\n")
    body_start = 0
    skip_kw = [
        "miroslav", "weberlis", "gerlikon", "8500", "079 602",
        "bewerbung als", "betreff", "@gmail", "z.hd", "im weberlis",
        "eichmatt", "rebberg",
    ]
    for idx, line in enumerate(lines):
        s = line.strip()
        if s.lower().startswith("sehr geehrte") or s.lower().startswith("dear"):
            body_start = idx
            break
        if any(kw in s.lower() for kw in skip_kw):
            continue
        if s and len(s) < 60 and not s.endswith(".") and not s.endswith(":"):
            continue
        if s:
            body_start = idx
            break
    clean_lines = lines[body_start:]

    # Separate: greeting, intro, "Was kann ich...", bullets, closing parts
    greeting = ""
    intro_paragraphs = []
    was_line = ""
    bullets = []
    closing_paragraphs = []

    section = "greeting"  # greeting → intro → bullets → closing
    for line in clean_lines:
        s = line.strip()
        if not s:
            continue

        if section == "greeting":
            if s.lower().startswith("sehr geehrte") or s.lower().startswith("dear"):
                greeting = s
                section = "intro"
            continue

        if section == "intro":
            if s.startswith("Was kann ich") or s.startswith("Was ich"):
                was_line = s
                section = "bullets"
            elif s.startswith("• ") or s.startswith("- ") or s.startswith("– "):
                was_line = "Was kann ich Ihrem Unternehmen bieten:"
                bullets.append(s.lstrip("•-– ").strip())
                section = "bullets"
            else:
                intro_paragraphs.append(s)
            continue

        if section == "bullets":
            if s.startswith("• ") or s.startswith("- ") or s.startswith("– "):
                bullets.append(s.lstrip("•-– ").strip())
            elif s:
                # First non-bullet after bullets → closing
                closing_paragraphs.append(s)
                section = "closing"
            continue

        if section == "closing":
            if s == "Miroslav Mikulic":
                continue  # skip — already in template
            if s.lower().startswith("freundliche gr") or s.lower().startswith("mit freundlichen"):
                continue  # skip — already in template
            closing_paragraphs.append(s)

    if not greeting:
        greeting = "Sehr geehrte Damen und Herren"

    # ── Modify the DOCX template ──
    doc = Document(str(template_path))
    paras = doc.paragraphs

    def set_para_text(p, new_text):
        """Replace paragraph text preserving formatting of first run."""
        if not p.runs:
            p.text = new_text
            return
        # Keep first run's formatting, clear rest
        first_run = p.runs[0]
        for run in p.runs[1:]:
            run.text = ""
        first_run.text = new_text

    def _force_arial_11(p):
        """Force all runs in paragraph to Arial 11pt."""
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    def set_tab_para(p, right_text):
        """Set text for a paragraph that uses [TAB]right_text format."""
        if not p.runs:
            p.text = f"\t{right_text}"
            return
        # First run should be tab, second+ should be text
        if len(p.runs) >= 2:
            p.runs[0].text = "\t"
            p.runs[1].text = right_text
            for run in p.runs[2:]:
                run.text = ""
        else:
            p.runs[0].text = f"\t{right_text}"

    # P0: Name (unchanged)
    # P1: Street
    set_para_text(paras[1], "Im Weberlis Rebberg 42")
    # P2: City
    set_para_text(paras[2], "8500 Gerlikon")
    # P3: Phone [TAB] Company
    if len(paras[3].runs) >= 2:
        paras[3].runs[0].text = "079 602 83 31"
        paras[3].runs[1].text = f"\t{clean_company}" if clean_company else "\t"
        for run in paras[3].runs[2:]:
            run.text = ""
    else:
        set_para_text(paras[3], f"079 602 83 31\t{clean_company or ''}")

    # P4-P5: Company address lines (right-aligned via tab)
    # Split address into parts: e.g. "Musterstr. 10\n8005 Zürich" → ["Musterstr. 10", "8005 Zürich"]
    addr_parts = []
    if company_address:
        addr_parts = [l.strip() for l in company_address.strip().split("\n") if l.strip()]

    if contact_person and len(addr_parts) >= 2:
        # Contact person + 2 address lines = 3 items, but only 2 slots (P4, P5)
        # Combine: P4 = contact person, P5 = "Strasse Nr, PLZ Ort"
        set_tab_para(paras[4], contact_person)
        set_tab_para(paras[5], ", ".join(addr_parts))
    elif contact_person and len(addr_parts) == 1:
        set_tab_para(paras[4], contact_person)
        set_tab_para(paras[5], addr_parts[0])
    elif contact_person:
        set_tab_para(paras[4], contact_person)
        set_tab_para(paras[5], "")
    elif len(addr_parts) >= 2:
        # No contact person: P4 = Strasse, P5 = PLZ Ort
        set_tab_para(paras[4], addr_parts[0])
        set_tab_para(paras[5], addr_parts[1])
    elif len(addr_parts) == 1:
        set_tab_para(paras[4], "")
        set_tab_para(paras[5], addr_parts[0])
    else:
        set_tab_para(paras[4], "")
        set_tab_para(paras[5], "")

    # P7: [TAB] Date
    date_str = f"Gerlikon, {datetime.now().strftime('%d.%m.%Y')}"
    set_tab_para(paras[7], date_str)

    # P10: Betreff (bold)
    betreff = job_title.strip() if job_title else "Bewerbung"
    set_para_text(paras[10], betreff)

    # P12: Greeting
    set_para_text(paras[12], greeting)

    # P14: Intro text — join all intro paragraphs
    intro_text = " ".join(intro_paragraphs) if intro_paragraphs else paras[14].text
    set_para_text(paras[14], intro_text)

    # P16: "Was kann ich..."
    set_para_text(paras[16], was_line or "Was kann ich Ihrem Unternehmen bieten:")

    def hide_para(p):
        """Make a paragraph virtually invisible (1pt font, no spacing)."""
        from docx.oxml.ns import qn
        from docx.shared import Pt as DPt
        set_para_text(p, "")
        # Remove bullet numbering
        pPr = p._element.pPr
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
        # Set font size to 1pt and space_after to 0
        if p.runs:
            p.runs[0].font.size = DPt(1)
        p.paragraph_format.space_after = DPt(0)
        p.paragraph_format.space_before = DPt(0)
        p.paragraph_format.line_spacing = 1.0

    # P18-P23: Bullets (template has 6 slots)
    bullet_slots = [18, 19, 20, 21, 22, 23]
    for i, slot in enumerate(bullet_slots):
        if i < len(bullets):
            set_para_text(paras[slot], bullets[i])
        else:
            hide_para(paras[slot])

    # P24: Empty bullet line in template → hide it
    hide_para(paras[24])

    # P25: Empty line after bullets → hide (save space)
    hide_para(paras[25])

    # P15: Empty line between intro and "Was kann ich" → hide to save space
    hide_para(paras[15])

    # P17: Empty line between "Was kann ich" and first bullet → hide
    hide_para(paras[17])

    # P13: Empty line after greeting → hide (tiny gap remains from space_after)
    hide_para(paras[13])

    # P27: Empty line before "Ich freue mich" → hide
    hide_para(paras[27])

    # Header empty lines: reduce P6, P8, P9, P11 to half-height
    # They still need SOME height for visual separation, but less
    from docx.shared import Pt as DPt
    for idx in [6, 8, 9, 11]:
        paras[idx].paragraph_format.space_after = DPt(2)
        paras[idx].paragraph_format.space_before = DPt(0)
        if paras[idx].runs:
            paras[idx].runs[0].font.size = DPt(6)

    # P26: First closing paragraph
    if closing_paragraphs:
        set_para_text(paras[26], closing_paragraphs[0])
    # P27: Empty line → keep as is (spacing between closing paragraphs)
    # P28: Second closing paragraph (or "Ich freue mich auf Ihren Anruf.")
    if len(closing_paragraphs) > 1:
        set_para_text(paras[28], closing_paragraphs[1])
    elif not closing_paragraphs:
        # Keep template default
        pass

    # P29: Empty line before signature → keep
    # P30: "Freundliche Grüsse" — keep as is
    # P31: "Miroslav Mikulic" — keep as is

    # ── Force Arial 11pt on all paragraphs ──
    for p in paras:
        _force_arial_11(p)

    # ── Save modified DOCX to temp file ──
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "letter.docx")
        doc.save(docx_path)

        # ── Convert to PDF via LibreOffice ──
        try:
            subprocess.run(
                [
                    "libreoffice", "--headless", "--norestore",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                capture_output=True, timeout=30, check=True,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
            print(f"[PDF] LibreOffice conversion failed: {e}")
            # Fallback to reportlab
            return _generate_cover_letter_pdf_fallback(
                letter_text, job_title, company, contact_person, company_address
            )

        pdf_path = os.path.join(tmpdir, "letter.pdf")
        if not os.path.exists(pdf_path):
            print("[PDF] LibreOffice produced no PDF output")
            return _generate_cover_letter_pdf_fallback(
                letter_text, job_title, company, contact_person, company_address
            )

        with open(pdf_path, "rb") as f:
            return f.read()


def _generate_cover_letter_pdf_fallback(
    letter_text: str,
    job_title: str = "",
    company: str = "",
    contact_person: str = "",
    company_address: str = "",
) -> bytes:
    """Reportlab fallback matching the DOCX template layout.

    Layout: sender left, recipient right (via tab at 9.75cm), date right,
    bullet points with hanging indent using • symbol.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    buffer = io.BytesIO()
    page_w, page_h = A4
    left_m = 2.5 * cm
    right_m = 2.5 * cm
    usable_w = page_w - left_m - right_m

    doc_pdf = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=2.5 * cm, bottomMargin=2.0 * cm,
        leftMargin=left_m, rightMargin=right_m,
    )

    # Register Arial if available on the system (Windows: C:\Windows\Fonts)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    _font_name = "Helvetica"
    _font_name_bold = "Helvetica-Bold"
    try:
        import platform
        if platform.system() == "Windows":
            pdfmetrics.registerFont(TTFont("Arial", r"C:\Windows\Fonts\arial.ttf"))
            pdfmetrics.registerFont(TTFont("Arial-Bold", r"C:\Windows\Fonts\arialbd.ttf"))
            _font_name = "Arial"
            _font_name_bold = "Arial-Bold"
        else:
            # Try common Linux paths
            for arial_path in ["/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
                               "/usr/share/fonts/truetype/arial.ttf"]:
                if os.path.exists(arial_path):
                    pdfmetrics.registerFont(TTFont("Arial", arial_path))
                    _font_name = "Arial"
                    break
    except Exception:
        pass  # Fallback to Helvetica

    # Styles
    s = ParagraphStyle("N", fontName=_font_name, fontSize=11, leading=14,
                       alignment=TA_LEFT, spaceAfter=6)
    s_right = ParagraphStyle("R", fontName=_font_name, fontSize=11, leading=14,
                             alignment=TA_RIGHT, spaceAfter=6)
    s_bold = ParagraphStyle("B", fontName=_font_name_bold, fontSize=11, leading=14,
                            alignment=TA_LEFT, spaceAfter=6)
    s_bullet = ParagraphStyle("BUL", fontName=_font_name, fontSize=11, leading=14,
                              alignment=TA_LEFT, spaceAfter=4,
                              leftIndent=14, firstLineIndent=-14)

    clean_company = _sanitize_company(company)

    # ── Header: Sender (left) | Recipient (right) ──
    sender_lines = [
        "Miroslav Mikulic",
        "Im Weberlis Rebberg 42",
        "8500 Gerlikon",
        "079 602 83 31",
    ]
    recipient_lines = []
    if clean_company:
        recipient_lines.append(clean_company)
    if contact_person:
        recipient_lines.append(contact_person)
    if company_address:
        for part in company_address.strip().split("\n"):
            if part.strip():
                recipient_lines.append(part.strip())

    # Pad to same length
    max_lines = max(len(sender_lines), len(recipient_lines))
    while len(sender_lines) < max_lines:
        sender_lines.append("")
    while len(recipient_lines) < max_lines:
        recipient_lines.append("")

    left_col_w = usable_w * 0.55
    right_col_w = usable_w * 0.45

    header_data = []
    for sl, rl in zip(sender_lines, recipient_lines):
        header_data.append([
            Paragraph(sl, s) if sl else Paragraph("", s),
            Paragraph(rl, s_right) if rl else Paragraph("", s_right),
        ])

    header_table = Table(header_data, colWidths=[left_col_w, right_col_w])
    header_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))

    story = [header_table, Spacer(1, 18)]

    # Date — right aligned
    date_str = f"Gerlikon, {datetime.now().strftime('%d.%m.%Y')}"
    story.append(Paragraph(date_str, s_right))
    story.append(Spacer(1, 24))

    # Betreff — bold
    story.append(Paragraph(job_title or "Bewerbung", s_bold))
    story.append(Spacer(1, 12))

    # Body text
    for line in letter_text.strip().split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
        elif stripped.startswith("- ") or stripped.startswith("• ") or stripped.startswith("– "):
            bullet_text = stripped.lstrip("-•– ").strip()
            story.append(Paragraph(f"•  {bullet_text}", s_bullet))
        elif stripped.startswith("Freundliche Gr") or stripped.startswith("Mit freundlichen"):
            story.append(Spacer(1, 4))
            story.append(Paragraph(stripped, s))
        elif stripped == "Miroslav Mikulic":
            story.append(Paragraph(stripped, s))
        else:
            story.append(Paragraph(stripped, s))

    doc_pdf.build(story)
    return buffer.getvalue()


def calculate_quick_score(job_title: str, job_company: str, job_location: str, cv_text: str) -> int:
    """Fast local keyword-based match score (0-100) without API calls.

    Scoring dimensions (max 100):
      - Role fit        (0-35): Does the job role match CV experience?
      - ERP/Tech fit    (0-25): Do specific systems/technologies match?
      - Seniority fit   (0-15): Is the level appropriate?
      - Domain fit      (0-15): Does the business domain match?
      - Location fit    (0-10): Is it in Deutschschweiz?
    """
    import re

    if not cv_text or not job_title:
        return 0

    cv_lower = cv_text.lower()
    title_lower = job_title.lower()
    score = 0

    # ═══ 1. ROLE FIT (max 35) ═══
    role_score = 0
    # Primary roles - only count the BEST match, not all
    primary_roles = [
        (["projektleiter", "project manager", "it-projektleiter", "projektmanager"], 35),
        (["business analyst", "anforderungsanalyst"], 30),
        (["consultant", "berater", "it consultant", "it-berater"], 25),
        (["entwickler", "developer", "engineer", "architekt"], 15),
        (["support", "helpdesk", "service desk"], 8),
        (["sachbearbeiter", "administration", "assistenz"], 5),
        (["verkauf", "sales", "account manager"], 5),
    ]
    for role_variants, points in primary_roles:
        if any(r in title_lower for r in role_variants):
            # Check if this role appears in CV
            if any(r in cv_lower for r in role_variants):
                role_score = max(role_score, points)
            else:
                role_score = max(role_score, int(points * 0.2))
            break  # Only match first (best) role category

    score += min(role_score, 35)

    # ═══ 2. ERP / TECHNOLOGY FIT (max 25) ═══
    tech_score = 0
    tech_matches = {
        # ERP systems (high value if specific match)
        "abacus": 15, "sap": 12, "microsoft dynamics": 12,
        "dynamics 365": 12, "business central": 12, "navision": 10,
        "oracle": 8, "salesforce": 8, "netsuite": 8, "proalpha": 8,
        # General tech
        "erp": 6, "crm": 5, "bi": 4, "sql": 3, "python": 3,
    }
    matched_tech = 0
    unmatched_tech = 0
    for tech, weight in tech_matches.items():
        if tech in title_lower:
            if tech in cv_lower:
                tech_score += weight
                matched_tech += 1
            else:
                unmatched_tech += 1
    # Penalty if job requires tech NOT in CV
    if unmatched_tech > 0 and matched_tech == 0:
        tech_score = max(0, tech_score - 5)
    score += min(tech_score, 25)

    # ═══ 3. SENIORITY FIT (max 15) ═══
    seniority_score = 10  # neutral default
    # Negative signals — junior roles
    if any(w in title_lower for w in ["junior", "trainee", "praktikant", "lehrling", "werkstudent", "stagiaire"]):
        seniority_score = 2
    # Positive — senior/lead aligns with Projektleiter experience
    elif any(w in title_lower for w in ["senior", "lead", "leiter", "head of", "director"]):
        if any(w in cv_lower for w in ["leiter", "projektleiter", "lead", "senior", "führung"]):
            seniority_score = 15
        else:
            seniority_score = 6
    score += seniority_score

    # ═══ 4. DOMAIN FIT (max 15) ═══
    domain_score = 0
    domains = {
        "finanzen": 4, "finance": 4, "rechnungswesen": 4, "controlling": 4,
        "logistik": 4, "supply chain": 4, "einkauf": 3, "produktion": 3,
        "personal": 3, "hr": 2, "fertigung": 3, "lager": 2,
        "digitalisierung": 5, "transformation": 4, "migration": 4,
        "implementierung": 4, "rollout": 4, "einführung": 4,
    }
    for domain, weight in domains.items():
        if domain in title_lower and domain in cv_lower:
            domain_score += weight
    score += min(domain_score, 15)

    # ═══ 5. LOCATION FIT (max 10) ═══
    deutschschweiz = [
        "zürich", "bern", "basel", "luzern", "winterthur", "st. gallen",
        "st.gallen", "aarau", "zug", "schaffhausen", "frauenfeld", "baden",
        "olten", "solothurn", "thun", "biel", "rapperswil", "wil", "chur",
        "gerlikon", "weinfelden", "kreuzlingen",
    ]
    romandie = ["genève", "lausanne", "fribourg", "neuchâtel", "sion", "nyon", "morges"]
    if job_location:
        loc_lower = job_location.lower()
        if any(c in loc_lower for c in deutschschweiz):
            score += 10
        elif any(c in loc_lower for c in romandie):
            score += 3  # French-speaking = less ideal
        elif "schweiz" in loc_lower or "remote" in loc_lower:
            score += 7

    # ═══ 6. LANGUAGE PENALTY ═══
    # French/Italian job titles are less relevant
    french_signals = ["responsable", "ingénieur", "chargé", "gestionnaire", "conseiller"]
    if any(f in title_lower for f in french_signals):
        score = int(score * 0.5)

    return max(0, min(100, int(score)))


def calculate_match_score(
    job_description: str,
    cv_text: str,
    api_key: str = ""
) -> dict:
    """Calculate how well a candidate matches a job posting."""

    client = _get_client()

    prompt = f"""Analysiere die Übereinstimmung zwischen diesem Stelleninserat und dem Lebenslauf.

**Stelleninserat:**
{job_description[:1500]}

**Lebenslauf:**
{cv_text[:1500]}

Antworte NUR mit einem JSON-Objekt in diesem Format:
{{
  "score": <Zahl 0-100>,
  "strengths": ["Stärke 1", "Stärke 2", "Stärke 3"],
  "gaps": ["Lücke 1", "Lücke 2"],
  "recommendation": "Kurze Empfehlung in 1-2 Sätzen"
}}"""

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=500,
        messages=[{"role": "user", "content": prompt}]
    )

    import json
    text = message.content[0].text.strip()
    # Extract JSON
    import re
    json_match = re.search(r'\{.*\}', text, re.DOTALL)
    if json_match:
        return json.loads(json_match.group(0))
    return {"score": 0, "strengths": [], "gaps": [], "recommendation": "Fehler bei der Analyse"}
